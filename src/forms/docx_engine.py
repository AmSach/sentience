"""
DOCX Engine for Sentience v3.0
Handles DOCX template filling, mail merge, table operations, and style preservation.
"""

import os
import re
import copy
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Iterator
from dataclasses import dataclass, field
from datetime import datetime, date
from string import Template

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


# Placeholder patterns
PLACEHOLDER_PATTERNS = {
    "curly_braces": r"\{\{(\w+)\}\}",
    "double_brackets": r"\[\[(\w+)\]\]",
    "percent": r"\%\%(\w+)\%\%",
    "mustache": r"\{\{([^}]+)\}\}",
}


@dataclass
class DOCXField:
    """Represents a DOCX placeholder field."""
    name: str
    placeholder: str
    location: str = "text"  # text, header, footer, table, textbox
    paragraph_index: int = 0
    element_index: int = 0
    table_index: Optional[int] = None
    cell_index: Optional[int] = None
    field_type: str = "text"
    default_value: Any = None


@dataclass
class DOCXTableInfo:
    """Information about a table in the document."""
    index: int
    rows: int
    cols: int
    has_header: bool
    header_row: List[str]
    is_merge_template: bool = False
    merge_column: Optional[int] = None


class DOCXEngine:
    """
    DOCX document operations engine.
    
    Handles:
    - Template filling with placeholder replacement
    - Mail merge for bulk document generation
    - Table operations and data insertion
    - Style preservation during edits
    """
    
    def __init__(
        self,
        docx_path: Optional[Union[str, Path]] = None,
        placeholder_style: str = "curly_braces"
    ):
        """
        Initialize DOCX engine.
        
        Args:
            docx_path: Path to DOCX file.
            placeholder_style: Style of placeholders (curly_braces, double_brackets, percent).
        """
        self.docx_path = Path(docx_path) if docx_path else None
        self.document: Optional[Document] = None
        self.placeholder_style = placeholder_style
        self.placeholder_pattern = PLACEHOLDER_PATTERNS.get(
            placeholder_style, PLACEHOLDER_PATTERNS["curly_braces"]
        )
        self.fields: Dict[str, DOCXField] = {}
        self.tables: List[DOCXTableInfo] = []
        self._original_styles: Dict[str, Any] = {}
        
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX operations. Install with: pip install python-docx"
            )
        
        if self.docx_path and self.docx_path.exists():
            self._load_document()
    
    def _load_document(self) -> None:
        """Load DOCX file."""
        try:
            self.document = Document(str(self.docx_path))
            logger.info(f"Loaded DOCX: {self.docx_path}")
            self._preserve_styles()
        except Exception as e:
            logger.error(f"Failed to load DOCX: {e}")
            raise
    
    def _preserve_styles(self) -> None:
        """Preserve original document styles for later restoration."""
        if not self.document:
            return
        
        # Store style information
        for style in self.document.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                self._original_styles[style.name] = {
                    "font_name": style.font.name,
                    "font_size": style.font.size,
                    "font_bold": style.font.bold,
                    "font_italic": style.font.italic,
                    "font_color": style.font.color,
                }
    
    def create_document(self) -> None:
        """Create a new empty document."""
        self.document = Document()
        logger.info("Created new DOCX document")
    
    def detect_fields(self) -> Dict[str, DOCXField]:
        """
        Detect all placeholder fields in the document.
        
        Returns:
            Dictionary mapping field names to DOCXField objects.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        self.fields.clear()
        
        # Scan document body
        self._scan_paragraphs(self.document.paragraphs, "text")
        
        # Scan tables
        for table_idx, table in enumerate(self.document.tables):
            self._scan_table(table, table_idx)
        
        # Scan headers
        for section in self.document.sections:
            if section.header.paragraphs:
                self._scan_paragraphs(section.header.paragraphs, "header")
            if section.header.tables:
                for table_idx, table in enumerate(section.header.tables):
                    self._scan_table(table, table_idx, "header")
        
        # Scan footers
        for section in self.document.sections:
            if section.footer.paragraphs:
                self._scan_paragraphs(section.footer.paragraphs, "footer")
            if section.footer.tables:
                for table_idx, table in enumerate(section.footer.tables):
                    self._scan_table(table, table_idx, "footer")
        
        # Analyze tables
        self._analyze_tables()
        
        logger.info(f"Detected {len(self.fields)} placeholder fields")
        return self.fields
    
    def _scan_paragraphs(self, paragraphs: List, location: str) -> None:
        """Scan paragraphs for placeholders."""
        for para_idx, paragraph in enumerate(paragraphs):
            text = paragraph.text
            matches = re.finditer(self.placeholder_pattern, text)
            
            for match in matches:
                field_name = match.group(1).strip()
                if field_name not in self.fields:
                    self.fields[field_name] = DOCXField(
                        name=field_name,
                        placeholder=match.group(0),
                        location=location,
                        paragraph_index=para_idx
                    )
    
    def _scan_table(
        self,
        table: Any,
        table_idx: int,
        location: str = "table"
    ) -> None:
        """Scan table cells for placeholders."""
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text
                    matches = re.finditer(self.placeholder_pattern, text)
                    
                    for match in matches:
                        field_name = match.group(1).strip()
                        if field_name not in self.fields:
                            self.fields[field_name] = DOCXField(
                                name=field_name,
                                placeholder=match.group(0),
                                location=location,
                                paragraph_index=para_idx,
                                table_index=table_idx,
                                cell_index=cell_idx
                            )
    
    def _analyze_tables(self) -> None:
        """Analyze tables for structure and merge templates."""
        if not self.document:
            return
        
        self.tables.clear()
        
        for table_idx, table in enumerate(self.document.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            
            # Check first row for headers
            has_header = False
            header_row = []
            if rows > 0:
                header_row = [cell.text.strip() for cell in table.rows[0].cells]
                has_header = all(header_row)  # Non-empty headers
            
            # Check if it's a merge template (row with placeholders)
            is_merge_template = False
            merge_column = None
            
            if rows >= 2:
                for cell_idx, cell in enumerate(table.rows[1].cells):
                    if re.search(self.placeholder_pattern, cell.text):
                        is_merge_template = True
                        merge_column = cell_idx
                        break
            
            self.tables.append(DOCXTableInfo(
                index=table_idx,
                rows=rows,
                cols=cols,
                has_header=has_header,
                header_row=header_row,
                is_merge_template=is_merge_template,
                merge_column=merge_column
            ))
    
    def fill_template(
        self,
        values: Dict[str, Any],
        preserve_style: bool = True,
        missing_value: str = ""
    ) -> None:
        """
        Fill template placeholders with provided values.
        
        Args:
            values: Dictionary mapping field names to values.
            preserve_style: Whether to preserve text formatting.
            missing_value: Value to use for missing fields.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        # Fill body paragraphs
        self._fill_paragraphs(self.document.paragraphs, values, preserve_style, missing_value)
        
        # Fill tables
        for table in self.document.tables:
            self._fill_table(table, values, preserve_style, missing_value)
        
        # Fill headers
        for section in self.document.sections:
            self._fill_paragraphs(section.header.paragraphs, values, preserve_style, missing_value)
            for table in section.header.tables:
                self._fill_table(table, values, preserve_style, missing_value)
        
        # Fill footers
        for section in self.document.sections:
            self._fill_paragraphs(section.footer.paragraphs, values, preserve_style, missing_value)
            for table in section.footer.tables:
                self._fill_table(table, values, preserve_style, missing_value)
        
        logger.info(f"Filled {len(values)} field values")
    
    def _fill_paragraphs(
        self,
        paragraphs: List,
        values: Dict[str, Any],
        preserve_style: bool,
        missing_value: str
    ) -> None:
        """Fill placeholders in paragraphs."""
        for paragraph in paragraphs:
            self._fill_paragraph(paragraph, values, preserve_style, missing_value)
    
    def _fill_paragraph(
        self,
        paragraph: Any,
        values: Dict[str, Any],
        preserve_style: bool,
        missing_value: str
    ) -> None:
        """Fill placeholders in a single paragraph while preserving formatting."""
        if not paragraph.text:
            return
        
        # Find all placeholders in paragraph
        text = paragraph.text
        matches = list(re.finditer(self.placeholder_pattern, text))
        
        if not matches:
            return
        
        # Build runs with preserved formatting
        if preserve_style:
            self._fill_preserve_style(paragraph, values, missing_value)
        else:
            # Simple replacement
            for match in matches:
                field_name = match.group(1).strip()
                replacement = str(values.get(field_name, missing_value))
                text = text.replace(match.group(0), replacement)
            paragraph.text = text
    
    def _fill_preserve_style(
        self,
        paragraph: Any,
        values: Dict[str, Any],
        missing_value: str
    ) -> None:
        """Fill placeholders while preserving run formatting."""
        # Collect run formatting before modification
        run_formats = []
        for run in paragraph.runs:
            run_formats.append({
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size,
                "font_color": run.font.color.rgb if run.font.color.rgb else None,
            })
        
        # Get the original text and replace placeholders
        original_text = paragraph.text
        new_text = original_text
        
        for match in re.finditer(self.placeholder_pattern, original_text):
            field_name = match.group(1).strip()
            replacement = str(values.get(field_name, missing_value))
            new_text = new_text.replace(match.group(0), replacement)
        
        # Clear paragraph and rebuild with formatting
        if run_formats:
            # Apply first run's formatting to entire text
            first_format = run_formats[0]
            paragraph.clear()
            run = paragraph.add_run(new_text)
            
            if first_format["bold"] is not None:
                run.bold = first_format["bold"]
            if first_format["italic"] is not None:
                run.italic = first_format["italic"]
            if first_format["underline"] is not None:
                run.underline = first_format["underline"]
            if first_format["font_name"]:
                run.font.name = first_format["font_name"]
            if first_format["font_size"]:
                run.font.size = first_format["font_size"]
            if first_format["font_color"]:
                run.font.color.rgb = first_format["font_color"]
    
    def _fill_table(
        self,
        table: Any,
        values: Dict[str, Any],
        preserve_style: bool,
        missing_value: str
    ) -> None:
        """Fill placeholders in table cells."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._fill_paragraph(paragraph, values, preserve_style, missing_value)
    
    def mail_merge(
        self,
        data_rows: List[Dict[str, Any]],
        output_dir: Union[str, Path],
        filename_template: str = "document_{index}.docx",
        preserve_style: bool = True
    ) -> List[Path]:
        """
        Perform mail merge generating multiple documents.
        
        Args:
            data_rows: List of data dictionaries for each document.
            output_dir: Directory to save merged documents.
            filename_template: Template for output filenames.
            preserve_style: Whether to preserve text formatting.
            
        Returns:
            List of paths to generated documents.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        generated_files = []
        
        for idx, row_data in enumerate(data_rows, 1):
            # Create a copy of the document
            doc_copy = Document(str(self.docx_path)) if self.docx_path else Document()
            
            # Create a temporary engine for the copy
            temp_engine = DOCXEngine.__new__(DOCXEngine)
            temp_engine.document = doc_copy
            temp_engine.placeholder_pattern = self.placeholder_pattern
            temp_engine.fields = {}
            temp_engine.tables = []
            
            # Fill the copy
            temp_engine._fill_paragraphs(doc_copy.paragraphs, row_data, preserve_style, "")
            
            for table in doc_copy.tables:
                temp_engine._fill_table(table, row_data, preserve_style, "")
            
            for section in doc_copy.sections:
                temp_engine._fill_paragraphs(section.header.paragraphs, row_data, preserve_style, "")
                temp_engine._fill_paragraphs(section.footer.paragraphs, row_data, preserve_style, "")
            
            # Save the document
            filename = filename_template.format(
                index=idx,
                **{k: str(v).replace("/", "-").replace("\\", "-") for k, v in row_data.items()}
            )
            output_path = output_dir / filename
            doc_copy.save(str(output_path))
            generated_files.append(output_path)
        
        logger.info(f"Generated {len(generated_files)} documents via mail merge")
        return generated_files
    
    def fill_table_data(
        self,
        table_index: int,
        data: List[List[Any]],
        has_header: bool = True
    ) -> None:
        """
        Fill a table with data rows.
        
        Args:
            table_index: Index of the table to fill.
            data: 2D list of data to insert.
            has_header: Whether data includes a header row.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        if table_index >= len(self.document.tables):
            raise ValueError(f"Table index {table_index} out of range")
        
        table = self.document.tables[table_index]
        
        # Determine start row
        start_row = 1 if has_header and len(table.rows) > 0 else 0
        
        # Ensure we have enough rows
        existing_rows = len(table.rows)
        required_rows = start_row + len(data)
        
        if required_rows > existing_rows:
            # Add rows as needed
            for _ in range(required_rows - existing_rows):
                table.add_row()
        
        # Fill data
        for row_idx, row_data in enumerate(data):
            actual_row_idx = start_row + row_idx
            if actual_row_idx >= len(table.rows):
                break
            
            row = table.rows[actual_row_idx]
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = str(cell_value)
        
        logger.info(f"Filled table {table_index} with {len(data)} rows")
    
    def add_table(
        self,
        data: List[List[Any]],
        style: str = "Table Grid",
        header: bool = True
    ) -> int:
        """
        Add a new table to the document.
        
        Args:
            data: 2D list of table data.
            style: Table style name.
            header: Whether first row is header.
            
        Returns:
            Index of the added table.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        if not data:
            return -1
        
        rows = len(data)
        cols = max(len(row) for row in data) if data else 0
        
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = style
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = str(cell_value)
                    
                    # Style header row
                    if header and row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        logger.info(f"Added table with {rows} rows and {cols} columns")
        return len(self.document.tables) - 1
    
    def set_cell_value(
        self,
        table_index: int,
        row: int,
        col: int,
        value: Any
    ) -> None:
        """Set a specific cell value in a table."""
        if not self.document:
            raise ValueError("No document loaded")
        
        table = self.document.tables[table_index]
        if row < len(table.rows) and col < len(table.rows[row].cells):
            table.rows[row].cells[col].text = str(value)
    
    def get_table_data(self, table_index: int) -> List[List[str]]:
        """Get all data from a table as 2D list."""
        if not self.document:
            raise ValueError("No document loaded")
        
        table = self.document.tables[table_index]
        return [
            [cell.text for cell in row.cells]
            for row in table.rows
        ]
    
    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        bold: bool = False,
        italic: bool = False
    ) -> None:
        """Add a paragraph to the document."""
        if not self.document:
            raise ValueError("No document loaded")
        
        paragraph = self.document.add_paragraph(text, style=style)
        
        if bold or italic:
            for run in paragraph.runs:
                run.bold = bold
                run.italic = italic
    
    def add_heading(
        self,
        text: str,
        level: int = 1
    ) -> None:
        """Add a heading to the document."""
        if not self.document:
            raise ValueError("No document loaded")
        
        self.document.add_heading(text, level=level)
    
    def add_image(
        self,
        image_path: Union[str, Path],
        width: Optional[float] = None,
        height: Optional[float] = None
    ) -> None:
        """Add an image to the document."""
        if not self.document:
            raise ValueError("No document loaded")
        
        image_path = Path(image_path)
        if not image_path.exists():
            raise FileNotFoundError(f"Image not found: {image_path}")
        
        kwargs = {}
        if width:
            kwargs["width"] = Inches(width)
        if height:
            kwargs["height"] = Inches(height)
        
        self.document.add_picture(str(image_path), **kwargs)
    
    def replace_text(
        self,
        search: str,
        replace: str,
        match_case: bool = False,
        whole_words: bool = False
    ) -> int:
        """
        Replace all occurrences of text in the document.
        
        Returns:
            Number of replacements made.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        count = 0
        pattern = re.escape(search)
        
        if whole_words:
            pattern = r"\b" + pattern + r"\b"
        
        flags = 0 if match_case else re.IGNORECASE
        
        # Replace in paragraphs
        for paragraph in self.document.paragraphs:
            if re.search(pattern, paragraph.text, flags):
                new_text = re.sub(pattern, replace, paragraph.text, flags=flags)
                # Preserve style while replacing
                self._replace_preserve_style(paragraph, paragraph.text, new_text)
                count += 1
        
        # Replace in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if re.search(pattern, paragraph.text, flags):
                            new_text = re.sub(pattern, replace, paragraph.text, flags=flags)
                            self._replace_preserve_style(paragraph, paragraph.text, new_text)
                            count += 1
        
        logger.info(f"Replaced {count} occurrences of '{search}' with '{replace}'")
        return count
    
    def _replace_preserve_style(self, paragraph: Any, old_text: str, new_text: str) -> None:
        """Replace text while preserving style."""
        if not paragraph.runs:
            paragraph.text = new_text
            return
        
        # Get formatting from first run
        first_run = paragraph.runs[0]
        paragraph.clear()
        run = paragraph.add_run(new_text)
        
        if first_run.bold is not None:
            run.bold = first_run.bold
        if first_run.italic is not None:
            run.italic = first_run.italic
        if first_run.underline is not None:
            run.underline = first_run.underline
        if first_run.font.name:
            run.font.name = first_run.font.name
        if first_run.font.size:
            run.font.size = first_run.font.size
    
    def save(self, output_path: Union[str, Path]) -> bool:
        """
        Save the document to a file.
        
        Args:
            output_path: Path to save the document.
            
        Returns:
            True if saved successfully.
        """
        if not self.document:
            raise ValueError("No document loaded")
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            self.document.save(str(output_path))
            logger.info(f"Saved DOCX to: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to save DOCX: {e}")
            raise
    
    def get_bytes(self) -> bytes:
        """Get the document as bytes."""
        if not self.document:
            raise ValueError("No document loaded")
        
        import io
        buffer = io.BytesIO()
        self.document.save(buffer)
        return buffer.getvalue()
    
    def get_field_summary(self) -> Dict[str, Dict[str, Any]]:
        """Get a summary of all detected fields."""
        return {
            name: {
                "placeholder": field.placeholder,
                "location": field.location,
                "table_index": field.table_index,
                "field_type": field.field_type
            }
            for name, field in self.fields.items()
        }
    
    def close(self) -> None:
        """Clean up resources."""
        self.document = None
        self.fields.clear()
        self.tables.clear()


# Convenience functions
def fill_docx_template(
    docx_path: Union[str, Path],
    values: Dict[str, Any],
    output_path: Union[str, Path],
    preserve_style: bool = True
) -> Dict[str, Any]:
    """
    Convenience function to fill a DOCX template.
    
    Args:
        docx_path: Path to source DOCX.
        values: Dictionary of field values.
        output_path: Path to save filled DOCX.
        preserve_style: Whether to preserve formatting.
        
    Returns:
        Dictionary with field summary and status.
    """
    engine = DOCXEngine(docx_path)
    engine.detect_fields()
    engine.fill_template(values, preserve_style)
    engine.save(output_path)
    
    result = {
        "success": True,
        "output_path": str(output_path),
        "fields_filled": len(values),
        "field_summary": engine.get_field_summary()
    }
    
    engine.close()
    return result


def mail_merge_docs(
    template_path: Union[str, Path],
    data_rows: List[Dict[str, Any]],
    output_dir: Union[str, Path]
) -> List[Path]:
    """
    Convenience function for mail merge.
    
    Args:
        template_path: Path to template DOCX.
        data_rows: List of data dictionaries.
        output_dir: Directory for output documents.
        
    Returns:
        List of generated document paths.
    """
    engine = DOCXEngine(template_path)
    engine.detect_fields()
    files = engine.mail_merge(data_rows, output_dir)
    engine.close()
    return files


def get_docx_fields(docx_path: Union[str, Path]) -> Dict[str, DOCXField]:
    """Convenience function to get DOCX placeholder fields."""
    engine = DOCXEngine(docx_path)
    fields = engine.detect_fields()
    engine.close()
    return fields
