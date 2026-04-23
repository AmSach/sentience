#!/usr/bin/env python3
"""Forms Engine - PDF and DOCX form filling"""
import os
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from pathlib import Path
import json

try:
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import DictionaryObject, ArrayObject, NameObject, TextStringObject
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

@dataclass
class FormField:
    name: str
    type: str  # text, checkbox, radio, dropdown
    value: Any = None
    page: int = 0
    rect: tuple = (0, 0, 0, 0)
    required: bool = False
    options: List[str] = None

class PDFFormEngine:
    """PDF form detection and filling"""
    
    def __init__(self):
        if not HAS_PYPDF:
            raise RuntimeError("pypdf not installed. Run: pip install pypdf")
            
    def detect_fields(self, pdf_path: str) -> List[FormField]:
        """Detect form fields in PDF"""
        reader = PdfReader(pdf_path)
        fields = []
        
        # Get form fields
        if "/AcroForm" in reader.trailer["/Root"]:
            form = reader.trailer["/Root"]["/AcroForm"]
            if "/Fields" in form:
                for i, field in enumerate(form["/Fields"]):
                    field_obj = field.get_object()
                    
                    field_name = str(field_obj.get("/T", f"field_{i}"))
                    field_type = "text"
                    
                    # Determine field type
                    if "/FT" in field_obj:
                        ft = str(field_obj["/FT"])
                        if ft == "/Btn":
                            field_type = "checkbox"
                        elif ft == "/Tx":
                            field_type = "text"
                        elif ft == "/Ch":
                            field_type = "dropdown"
                            
                    # Get field value
                    current_value = field_obj.get("/V", None)
                    
                    # Get field options for dropdowns
                    options = None
                    if "/Opt" in field_obj:
                        options = [str(opt) for opt in field_obj["/Opt"]]
                        
                    fields.append(FormField(
                        name=field_name,
                        type=field_type,
                        value=current_value,
                        options=options
                    ))
                    
        return fields
        
    def fill_pdf(self, pdf_path: str, output_path: str, field_values: Dict[str, Any]) -> bool:
        """Fill PDF form with values"""
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Clone pages
        for page in reader.pages:
            writer.add_page(page)
            
        # Fill fields
        if "/AcroForm" in reader.trailer["/Root"]:
            form = reader.trailer["/Root"]["/AcroForm"]
            
            if "/Fields" in form:
                for field in form["/Fields"]:
                    field_obj = field.get_object()
                    field_name = str(field_obj.get("/T", ""))
                    
                    if field_name in field_values:
                        value = field_values[field_name]
                        
                        # Update field value
                        if isinstance(value, bool):
                            # Checkbox
                            field_obj[NameObject("/V")] = NameObject("/Yes" if value else "/Off")
                        else:
                            # Text field
                            field_obj[NameObject("/V")] = TextStringObject(str(value))
                            
        # Write output
        with open(output_path, 'wb') as f:
            writer.write(f)
            
        return True
        
    def flatten_pdf(self, pdf_path: str, output_path: str):
        """Flatten PDF (make form fields read-only)"""
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            # Create content stream that draws field values
            writer.add_page(page)
            
        # Remove AcroForm
        if "/AcroForm" in writer._root_object:
            del writer._root_object["/AcroForm"]
            
        with open(output_path, 'wb') as f:
            writer.write(f)


class DOCXFormEngine:
    """DOCX template filling"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
            
    def detect_placeholders(self, docx_path: str, pattern: str = "{{") -> List[str]:
        """Detect placeholders in DOCX"""
        doc = Document(docx_path)
        placeholders = set()
        
        # Check paragraphs
        for para in doc.paragraphs:
            text = para.text
            start = 0
            while True:
                idx = text.find(pattern, start)
                if idx == -1:
                    break
                end = text.find("}}", idx)
                if end != -1:
                    placeholder = text[idx+len(pattern):end].strip()
                    placeholders.add(placeholder)
                    start = end + 2
                else:
                    break
                    
        return list(placeholders)
        
    def fill_template(self, docx_path: str, output_path: str, values: Dict[str, str]) -> bool:
        """Fill DOCX template with values"""
        doc = Document(docx_path)
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            for key, value in values.items():
                placeholder = "{{" + key + "}}"
                if placeholder in para.text:
                    # Replace while preserving formatting
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in values.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in para.text:
                                for run in para.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)
                                        
        doc.save(output_path)
        return True
        
    def mail_merge(self, docx_path: str, output_dir: str, data: List[Dict[str, str]]) -> List[str]:
        """Perform mail merge - create multiple documents from template"""
        output_files = []
        
        for i, row in enumerate(data):
            output_path = os.path.join(output_dir, f"merged_{i+1}.docx")
            self.fill_template(docx_path, output_path, row)
            output_files.append(output_path)
            
        return output_files


class FormManager:
    """Manage PDF/DOCX forms"""
    
    def __init__(self, storage_dir: str = None):
        self.storage_dir = Path(storage_dir or Path.home() / ".sentience" / "forms")
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        self.pdf_engine = PDFFormEngine() if HAS_PYPDF else None
        self.docx_engine = DOCXFormEngine() if HAS_DOCX else None
        
    def analyze_form(self, file_path: str) -> Dict:
        """Analyze a form file"""
        path = Path(file_path)
        
        if path.suffix.lower() == '.pdf':
            if not self.pdf_engine:
                return {"error": "PDF engine not available"}
                
            fields = self.pdf_engine.detect_fields(file_path)
            return {
                "type": "pdf",
                "fields": [
                    {"name": f.name, "type": f.type, "required": f.required, "options": f.options}
                    for f in fields
                ]
            }
            
        elif path.suffix.lower() in ['.docx', '.doc']:
            if not self.docx_engine:
                return {"error": "DOCX engine not available"}
                
            placeholders = self.docx_engine.detect_placeholders(file_path)
            return {
                "type": "docx",
                "fields": [{"name": p, "type": "text"} for p in placeholders]
            }
            
        else:
            return {"error": f"Unsupported file type: {path.suffix}"}
            
    def fill_form(self, file_path: str, output_path: str, values: Dict[str, Any]) -> Dict:
        """Fill a form with values"""
        path = Path(file_path)
        
        if path.suffix.lower() == '.pdf':
            if not self.pdf_engine:
                return {"error": "PDF engine not available"}
                
            success = self.pdf_engine.fill_pdf(file_path, output_path, values)
            return {"status": "success" if success else "error", "output": output_path}
            
        elif path.suffix.lower() in ['.docx', '.doc']:
            if not self.docx_engine:
                return {"error": "DOCX engine not available"}
                
            success = self.docx_engine.fill_template(file_path, output_path, values)
            return {"status": "success" if success else "error", "output": output_path}
            
        return {"error": "Unsupported file type"}
        
    def save_template(self, name: str, file_path: str, field_mapping: Dict = None):
        """Save a form template for reuse"""
        template_dir = self.storage_dir / name
        template_dir.mkdir(exist_ok=True)
        
        import shutil
        shutil.copy(file_path, template_dir / "template")
        
        if field_mapping:
            with open(template_dir / "mapping.json", 'w') as f:
                json.dump(field_mapping, f, indent=2)
                
    def load_template(self, name: str) -> Dict:
        """Load a saved template"""
        template_dir = self.storage_dir / name
        
        if not template_dir.exists():
            return {"error": f"Template {name} not found"}
            
        template_file = template_dir / "template"
        mapping_file = template_dir / "mapping.json"
        
        result = {"template": str(template_file)}
        
        if mapping_file.exists():
            with open(mapping_file) as f:
                result["mapping"] = json.load(f)
                
        return result


# Form tools for agent
class FormTools:
    """Agent tools for form operations"""
    
    def __init__(self, manager: FormManager):
        self.manager = manager
        
    def fill_pdf(self, pdf_path: str, output_path: str, field_values: Dict) -> Dict:
        """Fill PDF form with provided values"""
        return self.manager.fill_form(pdf_path, output_path, field_values)
        
    def fill_docx(self, docx_path: str, output_path: str, values: Dict) -> Dict:
        """Fill DOCX template with values"""
        return self.manager.fill_form(docx_path, output_path, values)
        
    def detect_fields(self, file_path: str) -> Dict:
        """Detect fields in a form"""
        return self.manager.analyze_form(file_path)
        
    def auto_fill_from_reference(self, form_path: str, output_path: str, 
                                  reference: Dict, field_mapping: Dict = None) -> Dict:
        """Auto-fill form from reference data (like another form or document)"""
        # Analyze form
        form_info = self.manager.analyze_form(form_path)
        
        if "error" in form_info:
            return form_info
            
        # Map reference data to fields
        values = {}
        for field in form_info.get("fields", []):
            field_name = field["name"]
            
            # Try to match by name
            for ref_key, ref_value in reference.items():
                if field_name.lower() in ref_key.lower() or ref_key.lower() in field_name.lower():
                    values[field_name] = ref_value
                    break
                    
        # Fill form
        return self.manager.fill_form(form_path, output_path, values)
