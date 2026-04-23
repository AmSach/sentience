"""Form engine - PDF, DOCX, XLSX processing and auto-fill."""
import json, io, re
from typing import Dict, List, Any, Optional

class FormField:
    def __init__(self, name: str, field_type: str, label: str = "", value: str = "", confidence: float = 0.0):
        self.name = name
        self.type = field_type  # text, number, date, checkbox, signature, dropdown
        self.label = label
        self.value = value
        self.confidence = confidence
        self.options: List[str] = []
        self.required = False
        self.read_only = False
        self.page = 0
        self.bounds: Dict[str, int] = {}

class Form:
    def __init__(self, name: str, form_type: str):
        self.name = name
        self.type = form_type  # pdf, docx, xlsx, web_form, image
        self.fields: List[FormField] = []
        self.pages: int = 0
        self.source_file: str = ""
        self.metadata: Dict[str, Any] = {}
        self.template_id: Optional[str] = None

class TemplateBuilder:
    """Build form templates for common use cases."""
    
    TEMPLATES = {
        "visa_application": {
            "name": "Visa Application",
            "fields": [
                FormField("full_name", "text", "Full Name", required=True),
                FormField("passport_number", "text", "Passport Number", required=True),
                FormField("date_of_birth", "date", "Date of Birth"),
                FormField("nationality", "dropdown", "Nationality"),
                FormField("visa_type", "dropdown", "Visa Type"),
                FormField("purpose_of_visit", "text", "Purpose of Visit"),
                FormField("entry_date", "date", "Expected Entry Date"),
                FormField("exit_date", "date", "Expected Exit Date"),
                FormField("accommodation_address", "text", "Accommodation Address"),
                FormField("signature", "signature", "Applicant Signature"),
            ]
        },
        "job_application": {
            "name": "Job Application",
            "fields": [
                FormField("full_name", "text", "Full Name", required=True),
                FormField("email", "text", "Email Address", required=True),
                FormField("phone", "text", "Phone Number"),
                FormField("current_company", "text", "Current Company"),
                FormField("experience_years", "number", "Years of Experience"),
                FormField("skills", "text", "Key Skills"),
                FormField("resume", "text", "Resume Link"),
                FormField("linkedin", "text", "LinkedIn Profile"),
                FormField("cover_letter", "text", "Cover Letter"),
                FormField("expected_salary", "number", "Expected Salary"),
                FormField("start_date", "date", "Available Start Date"),
                FormField("signature", "signature", "Applicant Signature"),
            ]
        },
        "government_form": {
            "name": "Government Form",
            "fields": [
                FormField("full_name", "text", "Full Name", required=True),
                FormField("father_name", "text", "Father's Name"),
                FormField("mother_name", "text", "Mother's Name"),
                FormField("address", "text", "Full Address"),
                FormField("aadhar", "text", "Aadhar Number"),
                FormField("pan", "text", "PAN Number"),
                FormField("date_of_birth", "date", "Date of Birth"),
                FormField("place_of_birth", "text", "Place of Birth"),
                FormField("gender", "dropdown", "Gender"),
                FormField("marital_status", "dropdown", "Marital Status"),
                FormField("occupation", "text", "Occupation"),
                FormField("annual_income", "number", "Annual Income"),
                FormField("signature", "signature", "Applicant Signature"),
                FormField("date", "date", "Date of Filing"),
            ]
        },
        "tax_form": {
            "name": "Tax Form",
            "fields": [
                FormField("full_name", "text", "Full Name", required=True),
                FormField("pan", "text", "PAN Number", required=True),
                FormField("aadhar", "text", "Aadhar Number"),
                FormField("income", "number", "Total Income", required=True),
                FormField("tax_exemptions", "number", "Tax Exemptions"),
                FormField("deductions_80c", "number", "Deductions 80C"),
                FormField("deductions_80d", "number", "Deductions 80D"),
                FormField("house_property_income", "number", "House Property Income"),
                FormField("other_income", "number", "Other Income"),
                FormField("tax_slab", "dropdown", "Tax Slab"),
                FormField("bank_account", "text", "Bank Account Number"),
                FormField("ifsc_code", "text", "IFSC Code"),
                FormField("signature", "signature", "Applicant Signature"),
                FormField("date", "date", "Date"),
            ]
        },
        "bank_account": {
            "name": "Bank Account Opening",
            "fields": [
                FormField("full_name", "text", "Full Name", required=True),
                FormField("date_of_birth", "date", "Date of Birth", required=True),
                FormField("address", "text", "Full Address", required=True),
                FormField("phone", "text", "Phone Number", required=True),
                FormField("email", "text", "Email Address"),
                FormField("aadhar", "text", "Aadhar Number"),
                FormField("pan", "text", "PAN Number"),
                FormField("occupation", "dropdown", "Occupation"),
                FormField("annual_income", "number", "Annual Income"),
                FormField("account_type", "dropdown", "Account Type"),
                FormField("initial_deposit", "number", "Initial Deposit"),
                FormField("nominee_name", "text", "Nominee Name"),
                FormField("nominee_relation", "dropdown", "Relation with Nominee"),
                FormField("signature", "signature", "Signature"),
                FormField("photo", "text", "Photo Upload"),
                FormField("date", "date", "Date"),
            ]
        },
        "insurance_claim": {
            "name": "Insurance Claim",
            "fields": [
                FormField("policy_number", "text", "Policy Number", required=True),
                FormField("claim_type", "dropdown", "Claim Type", required=True),
                FormField("incident_date", "date", "Incident Date", required=True),
                FormField("incident_description", "text", "Description of Incident"),
                FormField("hospital_name", "text", "Hospital Name"),
                FormField("treatment_start", "date", "Treatment Start Date"),
                FormField("treatment_end", "date", "Treatment End Date"),
                FormField("total_amount", "number", "Total Claim Amount"),
                FormField("bank_account", "text", "Bank Account for Payment"),
                FormField("documents_attached", "checkbox", "Documents Attached"),
                FormField("claimant_name", "text", "Claimant Name", required=True),
                FormField("claimant_phone", "text", "Claimant Phone"),
                FormField("signature", "signature", "Claimant Signature"),
                FormField("date", "date", "Date of Claim"),
            ]
        },
        "leave_application": {
            "name": "Leave Application",
            "fields": [
                FormField("employee_name", "text", "Employee Name", required=True),
                FormField("employee_id", "text", "Employee ID"),
                FormField("department", "text", "Department"),
                FormField("leave_type", "dropdown", "Leave Type", required=True),
                FormField("start_date", "date", "Start Date", required=True),
                FormField("end_date", "date", "End Date", required=True),
                FormField("total_days", "number", "Total Days"),
                FormField("reason", "text", "Reason for Leave"),
                FormField("emergency_contact", "text", "Emergency Contact"),
                FormField("work_handover", "text", "Work Handover Details"),
                FormField("signature", "signature", "Employee Signature"),
                FormField("date", "date", "Date"),
            ]
        },
        "contract_agreement": {
            "name": "Contract Agreement",
            "fields": [
                FormField("party_a_name", "text", "Party A Name", required=True),
                FormField("party_a_address", "text", "Party A Address"),
                FormField("party_b_name", "text", "Party B Name", required=True),
                FormField("party_b_address", "text", "Party B Address"),
                FormField("contract_date", "date", "Contract Date"),
                FormField("contract_value", "number", "Contract Value"),
                FormField("payment_terms", "text", "Payment Terms"),
                FormField("duration", "text", "Contract Duration"),
                FormField("terms_and_conditions", "text", "Terms and Conditions"),
                FormField("termination_clause", "text", "Termination Clause"),
                FormField("confidentiality", "checkbox", "Confidentiality Agreed"),
                FormField("governing_law", "dropdown", "Governing Law"),
                FormField("party_a_signature", "signature", "Party A Signature"),
                FormField("party_b_signature", "signature", "Party B Signature"),
                FormField("witness_name", "text", "Witness Name"),
                FormField("witness_signature", "signature", "Witness Signature"),
            ]
        }
    }
    
    @classmethod
    def list_templates(cls) -> List[str]:
        return list(cls.TEMPLATES.keys())
    
    @classmethod
    def get_template(cls, name: str) -> Optional[Dict]:
        return cls.TEMPLATES.get(name)
    
    @classmethod
    def create_form_from_template(cls, template_name: str) -> Form:
        tmpl = cls.TEMPLATES.get(template_name)
        if not tmpl:
            raise ValueError(f"Template '{template_name}' not found")
        form = Form(tmpl["name"], template_name)
        for f in tmpl["fields"]:
            form.fields.append(FormField(f["name"], f["type"], f.get("label",""), required=f.get("required", False)))
        return form

class FormProcessor:
    """Process PDF, DOCX, XLSX forms and extract/fill fields."""
    
    @staticmethod
    def extract_fields_from_pdf(pdf_path: str) -> Form:
        """Extract form fields from a PDF."""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(pdf_path)
            form = Form(os.path.basename(pdf_path), "pdf")
            form.pages = len(reader.pages)
            form.source_file = pdf_path
            
            for page in reader.pages:
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        obj = annot.get_object()
                        if obj.get("/Subtype") == "/Widget":
                            field_name = obj.get("/T", "unknown")
                            field_type = "text"
                            if obj.get("/FT") == "/Btn": field_type = "checkbox"
                            elif obj.get("/FT") == "/Ch": field_type = "dropdown"
                            form.fields.append(FormField(field_name, field_type, page=reader.pages.index(page)))
            return form
        except Exception as e:
            return Form("error", "pdf")
    
    @staticmethod
    def extract_fields_from_docx(docx_path: str) -> Form:
        """Extract structured fields from a DOCX document."""
        try:
            from docx import Document
            doc = Document(docx_path)
            form = Form(os.path.basename(docx_path), "docx")
            form.pages = len(doc.paragraphs)
            
            # Look for labeled fields (e.g., "Name: _____")
            for para in doc.paragraphs:
                text = para.text.strip()
                if ":" in text and len(text) < 200:
                    label = text.split(":")[0].strip()
                    field_name = label.lower().replace(" ", "_")
                    form.fields.append(FormField(field_name, "text", label))
            return form
        except:
            return Form("error", "docx")
    
    @staticmethod
    def extract_fields_from_xlsx(xlsx_path: str) -> Form:
        """Extract fields from an XLSX spreadsheet form."""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(xlsx_path)
            form = Form(os.path.basename(xlsx_path), "xlsx")
            
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(min_row=1, max_row=50, max_col=10):
                    for cell in row:
                        val = str(cell.value or "").strip()
                        if val and ":" in val and len(val) < 100:
                            label = val.split(":")[0].strip()
                            form.fields.append(FormField(label.lower().replace(" ", "_"), "text", label))
            return form
        except:
            return Form("error", "xlsx")

class AutoFillEngine:
    """Auto-fill forms using reference documents and AI."""
    
    def __init__(self, agent):
        self.agent = agent
    
    def fill_form(self, form: Form, reference_docs: List[str], user_profile: Dict) -> Form:
        """Fill form fields using reference docs + user profile."""
        for field in form.fields:
            # 1. Try user profile
            if field.name in user_profile:
                field.value = str(user_profile[field.name])
                field.confidence = 1.0
                continue
            
            # 2. Try reference documents
            best_match = self._search_reference(field, reference_docs)
            if best_match:
                field.value = best_match["value"]
                field.confidence = best_match["confidence"]
            
            # 3. Ask agent to infer
            if field.confidence < 0.5 and field.value:
                inferred = self._infer_field(field, user_profile)
                if inferred:
                    field.value = inferred["value"]
                    field.confidence = min(inferred["confidence"], 0.7)
        
        return form
    
    def _search_reference(self, field: FormField, docs: List[str]) -> Optional[Dict]:
        """Search reference documents for field value."""
        for doc in docs:
            try:
                if doc.endswith(".pdf"):
                    text = self._extract_pdf_text(doc)
                elif doc.endswith(".docx"):
                    text = self._extract_docx_text(doc)
                else:
                    text = open(doc).read()
                
                # Simple keyword match
                label_lower = field.label.lower()
                if label_lower in text.lower():
                    # Extract value near the label
                    idx = text.lower().find(label_lower)
                    snippet = text[idx:idx+200]
                    # Try to extract a meaningful value
                    match = re.search(r'[:—]\s*([A-Z0-9\s]{2,50})', snippet)
                    if match:
                        return {"value": match.group(1).strip(), "confidence": 0.8}
            except:
                pass
        return None
    
    def _infer_field(self, field: FormField, profile: Dict) -> Optional[Dict]:
        """Use AI to infer field value from context."""
        return None
    
    def _extract_pdf_text(self, path: str) -> str:
        try:
            from PyPDF2 import PdfReader
            return " ".join([p.extract_text() for p in PdfReader(path).pages])
        except:
            return ""
    
    def _extract_docx_text(self, path: str) -> str:
        try:
            from docx import Document
            return " ".join([p.text for p in Document(path).paragraphs])
        except:
            return ""

def create_form_from_template(template_name: str) -> Form:
    return TemplateBuilder.create_form_from_template(template_name)

def list_available_templates() -> List[str]:
    return TemplateBuilder.list_templates()

def extract_form_fields(file_path: str) -> Form:
    if file_path.endswith(".pdf"):
        return FormProcessor.extract_fields_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return FormProcessor.extract_fields_from_docx(file_path)
    elif file_path.endswith(".xlsx"):
        return FormProcessor.extract_fields_from_xlsx(file_path)
    return Form("unknown", "unknown")
